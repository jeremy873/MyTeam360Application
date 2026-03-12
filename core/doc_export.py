# ═══════════════════════════════════════════════════════════════════
# MyTeam360 — AI Platform
# Copyright © 2026 Praxis Holdings LLC. All rights reserved.
#
# PROPRIETARY AND CONFIDENTIAL — TRADE SECRET
# See LICENSE and NOTICE files for full legal terms.
# ═══════════════════════════════════════════════════════════════════

"""
Document Export Engine — Formal .docx generation with company letterhead.

Exports meeting minutes, roundtable transcripts, corporate records,
resolutions, and summaries as professional Word documents with:
  - Custom company letterhead (logo + name + address)
  - Consistent branding (fonts, colors, layout)
  - Headers and footers with page numbers
  - Proper formatting (headings, tables, bullet points)
  - Confidentiality notice
  - Digital timestamp for audit trail
"""

import os
import io
import json
import logging
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from .database import get_db

logger = logging.getLogger("MyTeam360.doc_export")

# Brand colors
GOLD = RGBColor(0xE8, 0xA3, 0x17)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x66, 0x66, 0x88)
LIGHT_GRAY = RGBColor(0xAA, 0xAA, 0xCC)


class LetterheadConfig:
    """Stores company letterhead settings."""

    def __init__(self):
        self.company_name = ""
        self.address_line1 = ""
        self.address_line2 = ""
        self.phone = ""
        self.email = ""
        self.website = ""
        self.logo_path = ""
        self.tagline = ""
        self.confidential = True

    def load_from_db(self, owner_id: str):
        """Load letterhead config from branding settings."""
        with get_db() as db:
            rows = db.execute(
                "SELECT key, value FROM branding").fetchall()
            settings = {r["key"]: r["value"] for r in rows}

            # Also check workspace settings
            ws_rows = db.execute(
                "SELECT key, value FROM workspace_settings").fetchall()
            ws = {r["key"]: r["value"] for r in ws_rows}

        self.company_name = settings.get("company_name", ws.get("workspace_name", "MyTeam360"))
        self.logo_path = settings.get("logo_path", "")
        self.email = settings.get("contact_email", "")
        self.website = settings.get("website", "")
        self.address_line1 = settings.get("address_line1", "")
        self.address_line2 = settings.get("address_line2", "")
        self.phone = settings.get("phone", "")
        self.tagline = settings.get("tagline", "")
        return self

    def load_from_dict(self, data: dict):
        """Load from API request data."""
        for key in ("company_name", "address_line1", "address_line2",
                     "phone", "email", "website", "logo_path", "tagline"):
            if key in data:
                setattr(self, key, data[key])
        self.confidential = data.get("confidential", True)
        return self


class DocExporter:
    """Generate professional .docx documents with letterhead."""

    def __init__(self):
        self.export_dir = os.path.join(
            os.path.dirname(os.path.dirname(__file__)), "data", "exports")
        os.makedirs(self.export_dir, exist_ok=True)

    def _create_doc(self, letterhead: LetterheadConfig) -> Document:
        """Create a new document with letterhead applied."""
        doc = Document()

        # Page setup
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1.2)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # Default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(11)
        font.color.rgb = DARK

        # Heading styles
        for level, size, color in [(1, 16, GOLD), (2, 13, DARK), (3, 11, DARK)]:
            hs = doc.styles[f"Heading {level}"]
            hs.font.name = "Arial"
            hs.font.size = Pt(size)
            hs.font.color.rgb = color
            hs.font.bold = True

        # Header — company letterhead
        header = section.header
        header.is_linked_to_previous = False

        if letterhead.company_name:
            hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Company name in gold
            run = hp.add_run(letterhead.company_name)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = GOLD
            run.font.name = "Arial"

            # Contact info
            if letterhead.address_line1 or letterhead.email:
                contact_parts = []
                if letterhead.address_line1:
                    contact_parts.append(letterhead.address_line1)
                if letterhead.address_line2:
                    contact_parts.append(letterhead.address_line2)
                if letterhead.phone:
                    contact_parts.append(letterhead.phone)
                if letterhead.email:
                    contact_parts.append(letterhead.email)
                if letterhead.website:
                    contact_parts.append(letterhead.website)

                cp = header.add_paragraph()
                cr = cp.add_run(" | ".join(contact_parts))
                cr.font.size = Pt(8)
                cr.font.color.rgb = GRAY
                cr.font.name = "Arial"

            # Divider line
            dp = header.add_paragraph()
            dp.paragraph_format.space_after = Pt(0)
            border_run = dp.add_run("─" * 85)
            border_run.font.size = Pt(6)
            border_run.font.color.rgb = GOLD

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if letterhead.confidential:
            cr = fp.add_run("CONFIDENTIAL")
            cr.font.size = Pt(7)
            cr.font.color.rgb = GRAY
            cr.font.bold = True
            cr.font.name = "Arial"
            fp.add_run("  |  ")

        fr = fp.add_run(f"Generated by {letterhead.company_name or 'MyTeam360'}")
        fr.font.size = Pt(7)
        fr.font.color.rgb = GRAY
        fr.font.name = "Arial"

        fp.add_run("  |  ")
        dr = fp.add_run(datetime.now().strftime("%B %d, %Y at %I:%M %p"))
        dr.font.size = Pt(7)
        dr.font.color.rgb = GRAY
        dr.font.name = "Arial"

        return doc

    def _add_title(self, doc: Document, title: str):
        """Add a document title."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(16)
        run = p.add_run(title)
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = DARK
        run.font.name = "Arial"

    def _add_meta_table(self, doc: Document, fields: dict):
        """Add a metadata table (Date, Participants, etc.)."""
        table = doc.add_table(rows=len(fields), cols=2)
        table.style = "Table Grid"
        for i, (key, value) in enumerate(fields.items()):
            row = table.rows[i]
            # Key cell
            cell_k = row.cells[0]
            cell_k.width = Inches(1.5)
            kp = cell_k.paragraphs[0]
            kr = kp.add_run(key)
            kr.font.size = Pt(9)
            kr.font.bold = True
            kr.font.color.rgb = DARK
            # Value cell
            cell_v = row.cells[1]
            cell_v.width = Inches(5)
            vp = cell_v.paragraphs[0]
            vr = vp.add_run(str(value))
            vr.font.size = Pt(9)
            vr.font.color.rgb = DARK
        doc.add_paragraph()  # spacing

    def _save(self, doc: Document, filename: str) -> str:
        """Save document and return file path."""
        if not filename.endswith(".docx"):
            filename += ".docx"
        filepath = os.path.join(self.export_dir, filename)
        doc.save(filepath)
        return filepath

    # ── EXPORT: Meeting Minutes ──────────────────────────────

    def export_minutes(self, minutes_data: dict,
                       letterhead: LetterheadConfig) -> str:
        """Export meeting minutes as .docx with letterhead."""
        doc = self._create_doc(letterhead)

        self._add_title(doc, "Meeting Minutes")

        participants = minutes_data.get("participants", [])
        if isinstance(participants, str):
            participants = json.loads(participants)

        self._add_meta_table(doc, {
            "Date": minutes_data.get("created_at", datetime.now().isoformat())[:10],
            "Subject": minutes_data.get("topic", ""),
            "Participants": ", ".join(participants) if participants else "N/A",
            "Status": minutes_data.get("status", "draft").upper(),
            "Document ID": minutes_data.get("id", ""),
        })

        # Minutes body
        minutes_text = minutes_data.get("minutes_text", "")
        for line in minutes_text.split("\n"):
            line = line.strip()
            if not line:
                continue
            # Detect headings (numbered sections or ALL CAPS)
            if (line[0].isdigit() and "." in line[:4]) or line.isupper():
                doc.add_heading(line, level=2)
            else:
                doc.add_paragraph(line)

        # Approval block
        if minutes_data.get("status") == "approved":
            doc.add_paragraph()
            ap = doc.add_paragraph()
            ar = ap.add_run("✓ APPROVED")
            ar.font.size = Pt(12)
            ar.font.bold = True
            ar.font.color.rgb = RGBColor(0x22, 0xC5, 0x5E)
            if minutes_data.get("approved_at"):
                doc.add_paragraph(f"Approved: {minutes_data['approved_at']}")
        else:
            doc.add_paragraph()
            dp = doc.add_paragraph()
            dr = dp.add_run("DRAFT — Pending Approval")
            dr.font.size = Pt(10)
            dr.font.italic = True
            dr.font.color.rgb = GRAY

        filename = f"minutes_{minutes_data.get('id', 'export')}_{datetime.now().strftime('%Y%m%d')}.docx"
        return self._save(doc, filename)

    # ── EXPORT: Roundtable Transcript ────────────────────────

    def export_roundtable(self, roundtable_data: dict,
                          letterhead: LetterheadConfig) -> str:
        """Export full Roundtable transcript as .docx."""
        doc = self._create_doc(letterhead)

        self._add_title(doc, "Roundtable Discussion Transcript")

        participants = roundtable_data.get("participants", [])
        if isinstance(participants, str):
            participants = json.loads(participants)
        names = [p.get("name", "Unknown") for p in participants]

        transcript = roundtable_data.get("transcript", [])
        if isinstance(transcript, str):
            transcript = json.loads(transcript)

        rounds_completed = sum(1 for t in transcript if t.get("type") == "round_marker")

        self._add_meta_table(doc, {
            "Date": roundtable_data.get("created_at", "")[:10],
            "Topic": roundtable_data.get("topic", ""),
            "Mode": roundtable_data.get("mode", "").replace("_", " ").title(),
            "Participants": ", ".join(names),
            "Rounds": str(rounds_completed),
            "Status": roundtable_data.get("status", "").upper(),
        })

        # Transcript
        doc.add_heading("Discussion Transcript", level=2)
        for entry in transcript:
            if entry.get("type") == "round_marker":
                rp = doc.add_paragraph()
                rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rr = rp.add_run(f"── Round {entry.get('round', '?')} Complete ──")
                rr.font.size = Pt(9)
                rr.font.color.rgb = GOLD
                rr.font.italic = True
            elif entry.get("type") in ("user", "agent", "human"):
                p = doc.add_paragraph()
                # Speaker name in bold gold
                name_run = p.add_run(f"{entry.get('name', 'Unknown')}: ")
                name_run.font.bold = True
                name_run.font.color.rgb = GOLD if entry.get("type") == "agent" else DARK
                name_run.font.size = Pt(10)
                # Content
                content_run = p.add_run(entry.get("content", ""))
                content_run.font.size = Pt(10)
                content_run.font.color.rgb = DARK

        # Summary if available
        if roundtable_data.get("summary"):
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(roundtable_data["summary"])

        filename = f"roundtable_{roundtable_data.get('id', 'export')}_{datetime.now().strftime('%Y%m%d')}.docx"
        return self._save(doc, filename)

    # ── EXPORT: Corporate Record ─────────────────────────────

    def export_record(self, record_data: dict,
                      letterhead: LetterheadConfig) -> str:
        """Export a corporate record as .docx."""
        doc = self._create_doc(letterhead)

        record_type = record_data.get("record_type", "general").replace("_", " ").title()
        self._add_title(doc, f"Corporate Record — {record_type}")

        tags = record_data.get("tags", [])
        if isinstance(tags, str):
            tags = json.loads(tags)

        self._add_meta_table(doc, {
            "Date": record_data.get("created_at", "")[:10],
            "Title": record_data.get("title", ""),
            "Type": record_type,
            "Tags": ", ".join(tags) if tags else "None",
            "Record ID": record_data.get("id", ""),
            "Retention": f"{record_data.get('retention_days', 0)} days" if record_data.get("retention_days") else "Permanent",
        })

        doc.add_heading("Content", level=2)
        content = record_data.get("content", "")
        for line in content.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

        filename = f"record_{record_data.get('id', 'export')}_{datetime.now().strftime('%Y%m%d')}.docx"
        return self._save(doc, filename)

    # ── EXPORT: Resolution ───────────────────────────────────

    def export_resolution(self, resolution_data: dict,
                          letterhead: LetterheadConfig) -> str:
        """Export a resolution with vote record as .docx."""
        doc = self._create_doc(letterhead)

        self._add_title(doc, "Resolution")

        approvers = resolution_data.get("required_approvers", [])
        if isinstance(approvers, str):
            approvers = json.loads(approvers)

        self._add_meta_table(doc, {
            "Date": resolution_data.get("created_at", "")[:10],
            "Title": resolution_data.get("title", ""),
            "Threshold": resolution_data.get("threshold", "majority").title(),
            "Status": resolution_data.get("status", "pending").upper(),
            "Resolution ID": resolution_data.get("id", ""),
        })

        doc.add_heading("Description", level=2)
        doc.add_paragraph(resolution_data.get("description", ""))

        # Vote record
        votes = resolution_data.get("votes", [])
        if votes:
            doc.add_heading("Vote Record", level=2)
            vote_table = doc.add_table(rows=1, cols=4)
            vote_table.style = "Table Grid"
            headers = ["Voter", "Vote", "Comment", "Date"]
            for i, h in enumerate(headers):
                cell = vote_table.rows[0].cells[i]
                p = cell.paragraphs[0]
                r = p.add_run(h)
                r.font.bold = True
                r.font.size = Pt(9)

            for vote in votes:
                if isinstance(vote, str):
                    vote = json.loads(vote)
                row = vote_table.add_row()
                row.cells[0].text = vote.get("voter_name", "")
                vote_text = vote.get("vote", "")
                row.cells[1].text = vote_text.upper()
                row.cells[2].text = vote.get("comment", "")
                row.cells[3].text = (vote.get("created_at", ""))[:10]

                # Color the vote
                vp = row.cells[1].paragraphs[0]
                for run in vp.runs:
                    if "APPROVE" in vote_text.upper():
                        run.font.color.rgb = RGBColor(0x22, 0xC5, 0x5E)
                    elif "REJECT" in vote_text.upper():
                        run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)

        # Outcome
        if resolution_data.get("status") in ("approved", "rejected"):
            doc.add_paragraph()
            op = doc.add_paragraph()
            status = resolution_data["status"].upper()
            color = RGBColor(0x22, 0xC5, 0x5E) if status == "APPROVED" else RGBColor(0xEF, 0x44, 0x44)
            or_run = op.add_run(f"OUTCOME: {status}")
            or_run.font.size = Pt(14)
            or_run.font.bold = True
            or_run.font.color.rgb = color
            if resolution_data.get("decided_at"):
                doc.add_paragraph(f"Decided: {resolution_data['decided_at']}")

        # Signature lines
        doc.add_paragraph()
        doc.add_paragraph()
        for approver in (approvers or ["_________________________"]):
            sp = doc.add_paragraph()
            sp.add_run("________________________________").font.color.rgb = GRAY
            doc.add_paragraph(approver)
            dp = doc.add_paragraph("Date: ________________")
            dp.runs[0].font.color.rgb = GRAY
            doc.add_paragraph()

        filename = f"resolution_{resolution_data.get('id', 'export')}_{datetime.now().strftime('%Y%m%d')}.docx"
        return self._save(doc, filename)

    # ── EXPORT: Custom Summary ───────────────────────────────

    def export_summary(self, title: str, content: str,
                       metadata: dict, letterhead: LetterheadConfig) -> str:
        """Export any summary or report as .docx."""
        doc = self._create_doc(letterhead)
        self._add_title(doc, title)

        if metadata:
            self._add_meta_table(doc, metadata)

        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("- ") or line.startswith("• "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line[0].isdigit() and "." in line[:4]:
                doc.add_paragraph(line, style="List Number")
            else:
                doc.add_paragraph(line)

        filename = f"summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        return self._save(doc, filename)


# ══════════════════════════════════════════════════════════════
# LETTERHEAD MANAGEMENT
# ══════════════════════════════════════════════════════════════

class LetterheadManager:
    """Save and manage company letterhead settings."""

    def save_letterhead(self, owner_id: str, data: dict) -> dict:
        """Save letterhead configuration."""
        keys = ["company_name", "address_line1", "address_line2",
                "phone", "email", "website", "tagline"]
        with get_db() as db:
            for key in keys:
                if key in data:
                    db.execute(
                        "INSERT OR REPLACE INTO branding (key, value) VALUES (?,?)",
                        (key, data[key]))
        return {"success": True, "saved_fields": [k for k in keys if k in data]}

    def get_letterhead(self, owner_id: str) -> dict:
        """Get current letterhead configuration."""
        config = LetterheadConfig()
        config.load_from_db(owner_id)
        return {
            "company_name": config.company_name,
            "address_line1": config.address_line1,
            "address_line2": config.address_line2,
            "phone": config.phone,
            "email": config.email,
            "website": config.website,
            "tagline": config.tagline,
            "logo_path": config.logo_path,
        }
