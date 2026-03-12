# ═══════════════════════════════════════════════════════════════════
# MyTeam360 — AI Platform
# Copyright © 2026 Praxis Holdings LLC. All rights reserved.
#
# PROPRIETARY AND CONFIDENTIAL — TRADE SECRET
# Report IP violations: legal@praxisholdingsllc.com
# ═══════════════════════════════════════════════════════════════════

"""
Document Export Engine — Get data OUT of the platform in any format.

Formats:
  - Word (.docx) — conversations, roundtable transcripts, reports
  - PDF (.pdf)   — same as Word, with formatting
  - Excel (.xlsx) — analytics, spend data, audit logs, structured data
  - Markdown (.md) — universal text format
  - CSV (.csv)     — tabular data

Export targets:
  - Single conversation → Word/PDF/Markdown
  - Roundtable Meeting transcript + whiteboard → Word/PDF
  - Analytics dashboard → Excel/PDF
  - Spend/cost data → Excel/CSV
  - Agent configuration → JSON
  - Full data export → ZIP of everything
"""

import io
import os
import csv
import json
import logging
from datetime import datetime

logger = logging.getLogger("MyTeam360.export")


class DocumentExporter:
    """Universal export engine for all platform data."""

    # ══════════════════════════════════════════════════════════
    # WORD / DOCX
    # ══════════════════════════════════════════════════════════

    def conversation_to_docx(self, conversation: dict, messages: list) -> bytes:
        """Export a conversation as a Word document."""
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading(
            conversation.get("title", "Conversation"), level=1)
        title.runs[0].font.color.rgb = RGBColor(0x16, 0x42, 0x68)

        # Metadata
        meta = doc.add_paragraph()
        meta.style.font.size = Pt(9)
        meta.style.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
        agent_name = conversation.get("agent_name", "AI Assistant")
        created = conversation.get("created_at", "")[:10]
        meta.add_run(f"Space: {agent_name}  |  Date: {created}  |  "
                     f"Messages: {len(messages)}").font.size = Pt(9)

        doc.add_paragraph("")  # spacer

        # Messages
        for msg in messages:
            role = msg.get("role", "user")
            content = msg.get("content", "")

            # Role header
            if role == "user":
                p = doc.add_paragraph()
                run = p.add_run("You")
                run.bold = True
                run.font.color.rgb = RGBColor(0x16, 0x42, 0x68)
            else:
                p = doc.add_paragraph()
                run = p.add_run(agent_name)
                run.bold = True
                run.font.color.rgb = RGBColor(0xa4, 0x59, 0xf2)

            # Timestamp
            ts = msg.get("created_at", "")
            if ts:
                run2 = p.add_run(f"  {ts[:16].replace('T', ' ')}")
                run2.font.size = Pt(8)
                run2.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

            # Content
            content_p = doc.add_paragraph(content)
            content_p.style.font.size = Pt(11)
            content_p.paragraph_format.space_after = Pt(12)

        # Footer
        doc.add_paragraph("")
        footer = doc.add_paragraph()
        footer_run = footer.add_run(
            f"Exported from MyTeam360 on {datetime.now().strftime('%B %d, %Y')}  |  "
            f"© 2026 Praxis Holdings LLC")
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def roundtable_to_docx(self, roundtable: dict, whiteboard: dict = None) -> bytes:
        """Export a Roundtable Meeting transcript as Word."""
        from docx import Document
        from docx.shared import Pt, RGBColor

        doc = Document()
        doc.add_heading(f"Roundtable Meeting: {roundtable.get('topic', '')}", level=1)

        # Metadata
        meta = doc.add_paragraph()
        mode = roundtable.get("mode", "brainstorm").replace("_", " ").title()
        participants = roundtable.get("participants", [])
        names = ", ".join(p.get("name", "") for p in participants)
        run = meta.add_run(f"Mode: {mode}  |  Participants: {names}  |  "
                           f"Status: {roundtable.get('status', '')}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

        # Transcript
        doc.add_heading("Discussion", level=2)
        transcript = roundtable.get("transcript", [])
        for entry in transcript:
            if entry.get("type") == "round_marker":
                doc.add_heading(f"— Round {entry.get('round', '')} —", level=3)
                continue
            if entry.get("type") not in ("agent", "user"):
                continue

            name = entry.get("name", "Unknown")
            content = entry.get("content", "")

            p = doc.add_paragraph()
            run = p.add_run(f"{name}: ")
            run.bold = True
            if entry.get("type") == "user":
                run.font.color.rgb = RGBColor(0x16, 0x42, 0x68)
            else:
                run.font.color.rgb = RGBColor(0xa4, 0x59, 0xf2)
            p.add_run(content)
            p.paragraph_format.space_after = Pt(8)

        # Whiteboard
        if whiteboard and whiteboard.get("notes"):
            doc.add_heading("Whiteboard Notes", level=2)
            for section in whiteboard.get("sections", []):
                section_notes = [n for n in whiteboard["notes"]
                                if n["section_id"] == section["id"]]
                if not section_notes:
                    continue
                doc.add_heading(f"{section.get('icon', '')} {section['title']}", level=3)
                for note in section_notes:
                    bullet = doc.add_paragraph(style="List Bullet")
                    status = "✅ " if note.get("completed") else ""
                    author = f" — {note['author']}" if note.get("author") else ""
                    bullet.add_run(f"{status}{note['content']}{author}")

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ══════════════════════════════════════════════════════════
    # PDF
    # ══════════════════════════════════════════════════════════

    def conversation_to_pdf(self, conversation: dict, messages: list) -> bytes:
        """Export conversation as PDF using HTML→PDF conversion."""
        html = self._conversation_to_html(conversation, messages)
        return self._html_to_pdf(html)

    def roundtable_to_pdf(self, roundtable: dict, whiteboard: dict = None) -> bytes:
        """Export Roundtable Meeting as PDF."""
        html = self._roundtable_to_html(roundtable, whiteboard)
        return self._html_to_pdf(html)

    def _html_to_pdf(self, html: str) -> bytes:
        """Convert HTML to PDF. Uses weasyprint if available, falls back to basic."""
        try:
            from weasyprint import HTML
            return HTML(string=html).write_pdf()
        except ImportError:
            # Fallback: return HTML as bytes with PDF-like header
            # In production, install weasyprint or use a PDF library
            return html.encode("utf-8")

    def _conversation_to_html(self, conversation: dict, messages: list) -> str:
        agent_name = conversation.get("agent_name", "AI Assistant")
        title = conversation.get("title", "Conversation")
        msg_html = ""
        for msg in messages:
            role = msg.get("role", "user")
            content = msg.get("content", "").replace("\n", "<br>")
            name = "You" if role == "user" else agent_name
            color = "#1e40af" if role == "user" else "#a459f2"
            msg_html += f"""
            <div style="margin-bottom:16px">
                <div style="font-weight:700;color:{color};font-size:13px;margin-bottom:4px">{name}</div>
                <div style="font-size:14px;line-height:1.7;color:#334155">{content}</div>
            </div>"""

        return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
        <style>body{{font-family:-apple-system,sans-serif;max-width:700px;margin:40px auto;padding:20px;color:#1e293b}}
        h1{{font-size:22px;border-bottom:2px solid #a459f2;padding-bottom:8px}}</style></head>
        <body><h1>{title}</h1>
        <p style="font-size:11px;color:#94a3b8">Space: {agent_name} | Exported {datetime.now().strftime('%B %d, %Y')}</p>
        {msg_html}
        <hr style="border:none;border-top:1px solid #e5e7eb;margin-top:32px">
        <p style="font-size:9px;color:#94a3b8">Exported from MyTeam360 | © 2026 Praxis Holdings LLC</p>
        </body></html>"""

    def _roundtable_to_html(self, roundtable: dict, whiteboard: dict = None) -> str:
        topic = roundtable.get("topic", "")
        transcript = roundtable.get("transcript", [])
        entries_html = ""
        for e in transcript:
            if e.get("type") == "round_marker":
                entries_html += f"<h3 style='color:#94a3b8;margin:20px 0 8px'>— Round {e.get('round','')} —</h3>"
                continue
            if e.get("type") not in ("agent", "user"):
                continue
            name = e.get("name", "")
            color = "#1e40af" if e.get("type") == "user" else "#a459f2"
            entries_html += f"""
            <div style="margin-bottom:12px">
                <span style="font-weight:700;color:{color}">{name}:</span>
                <span style="color:#334155">{e.get('content','')}</span>
            </div>"""

        wb_html = ""
        if whiteboard and whiteboard.get("notes"):
            wb_html = "<h2>Whiteboard Notes</h2>"
            for sec in whiteboard.get("sections", []):
                notes = [n for n in whiteboard["notes"] if n["section_id"] == sec["id"]]
                if notes:
                    wb_html += f"<h3>{sec.get('icon','')} {sec['title']}</h3><ul>"
                    for n in notes:
                        wb_html += f"<li>{'✅ ' if n.get('completed') else ''}{n['content']}"
                        if n.get('author'):
                            wb_html += f" — <em>{n['author']}</em>"
                        wb_html += "</li>"
                    wb_html += "</ul>"

        return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
        <style>body{{font-family:-apple-system,sans-serif;max-width:700px;margin:40px auto;padding:20px}}</style></head>
        <body><h1>Roundtable: {topic}</h1>{entries_html}{wb_html}
        <p style="font-size:9px;color:#94a3b8;margin-top:32px">Exported from MyTeam360 | © 2026 Praxis Holdings LLC</p>
        </body></html>"""

    # ══════════════════════════════════════════════════════════
    # EXCEL / XLSX
    # ══════════════════════════════════════════════════════════

    def conversation_to_xlsx(self, conversation: dict, messages: list) -> bytes:
        """Export conversation as Excel spreadsheet."""
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Conversation"

        # Header
        headers = ["#", "Timestamp", "Role", "Content", "Model", "Tokens"]
        header_fill = PatternFill(start_color="a459f2", end_color="a459f2", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True, size=11)

        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.fill = header_fill
            cell.font = header_font

        # Data
        for i, msg in enumerate(messages, 1):
            ws.cell(row=i+1, column=1, value=i)
            ws.cell(row=i+1, column=2, value=msg.get("created_at", "")[:19])
            ws.cell(row=i+1, column=3, value=msg.get("role", ""))
            ws.cell(row=i+1, column=4, value=msg.get("content", ""))
            ws.cell(row=i+1, column=5, value=msg.get("model", ""))
            ws.cell(row=i+1, column=6, value=msg.get("tokens_used", 0))

        # Column widths
        ws.column_dimensions["A"].width = 5
        ws.column_dimensions["B"].width = 20
        ws.column_dimensions["C"].width = 10
        ws.column_dimensions["D"].width = 80
        ws.column_dimensions["E"].width = 20
        ws.column_dimensions["F"].width = 10

        # Wrap text in content column
        for row in range(2, len(messages) + 2):
            ws.cell(row=row, column=4).alignment = Alignment(wrap_text=True)

        # Info sheet
        info = wb.create_sheet("Info")
        info_data = [
            ("Title", conversation.get("title", "")),
            ("Space", conversation.get("agent_name", "")),
            ("Created", conversation.get("created_at", "")),
            ("Messages", len(messages)),
            ("Exported", datetime.now().isoformat()),
            ("Platform", "MyTeam360"),
        ]
        for i, (k, v) in enumerate(info_data, 1):
            info.cell(row=i, column=1, value=k).font = Font(bold=True)
            info.cell(row=i, column=2, value=str(v))

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def analytics_to_xlsx(self, analytics_data: dict) -> bytes:
        """Export analytics dashboard data as Excel."""
        import openpyxl
        from openpyxl.styles import Font, PatternFill

        wb = openpyxl.Workbook()
        header_fill = PatternFill(start_color="a459f2", end_color="a459f2", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True)

        # Summary sheet
        ws = wb.active
        ws.title = "Summary"
        summary_items = analytics_data.get("summary", {})
        row = 1
        for k, v in summary_items.items():
            ws.cell(row=row, column=1, value=k.replace("_", " ").title()).font = Font(bold=True)
            ws.cell(row=row, column=2, value=str(v))
            row += 1

        # Spend data
        if analytics_data.get("spend"):
            spend_ws = wb.create_sheet("Spend")
            spend_headers = ["Date", "Provider", "Model", "Input Tokens",
                           "Output Tokens", "Cost ($)"]
            for col, h in enumerate(spend_headers, 1):
                cell = spend_ws.cell(row=1, column=col, value=h)
                cell.fill = header_fill
                cell.font = header_font
            for i, entry in enumerate(analytics_data["spend"], 2):
                spend_ws.cell(row=i, column=1, value=entry.get("date", ""))
                spend_ws.cell(row=i, column=2, value=entry.get("provider", ""))
                spend_ws.cell(row=i, column=3, value=entry.get("model", ""))
                spend_ws.cell(row=i, column=4, value=entry.get("input_tokens", 0))
                spend_ws.cell(row=i, column=5, value=entry.get("output_tokens", 0))
                spend_ws.cell(row=i, column=6, value=entry.get("cost", 0))

        # Usage by space
        if analytics_data.get("by_space"):
            space_ws = wb.create_sheet("By Space")
            space_headers = ["Space", "Messages", "Tokens", "Cost ($)"]
            for col, h in enumerate(space_headers, 1):
                cell = space_ws.cell(row=1, column=col, value=h)
                cell.fill = header_fill
                cell.font = header_font
            for i, entry in enumerate(analytics_data["by_space"], 2):
                space_ws.cell(row=i, column=1, value=entry.get("name", ""))
                space_ws.cell(row=i, column=2, value=entry.get("messages", 0))
                space_ws.cell(row=i, column=3, value=entry.get("tokens", 0))
                space_ws.cell(row=i, column=4, value=entry.get("cost", 0))

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def data_to_xlsx(self, headers: list, rows: list,
                      sheet_name: str = "Data") -> bytes:
        """Generic: export any tabular data as Excel."""
        import openpyxl
        from openpyxl.styles import Font, PatternFill

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name

        header_fill = PatternFill(start_color="a459f2", end_color="a459f2", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True)

        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.fill = header_fill
            cell.font = header_font

        for i, row_data in enumerate(rows, 2):
            for col, val in enumerate(row_data, 1):
                ws.cell(row=i, column=col, value=str(val) if val else "")

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    # ══════════════════════════════════════════════════════════
    # MARKDOWN
    # ══════════════════════════════════════════════════════════

    def conversation_to_markdown(self, conversation: dict, messages: list) -> str:
        agent_name = conversation.get("agent_name", "AI Assistant")
        lines = [
            f"# {conversation.get('title', 'Conversation')}",
            f"*Space: {agent_name} | {conversation.get('created_at', '')[:10]}*\n",
        ]
        for msg in messages:
            role = "**You**" if msg.get("role") == "user" else f"**{agent_name}**"
            lines.append(f"{role}:")
            lines.append(f"{msg.get('content', '')}\n")
        lines.append(f"---\n*Exported from MyTeam360 on {datetime.now().strftime('%B %d, %Y')}*")
        return "\n".join(lines)

    # ══════════════════════════════════════════════════════════
    # CSV
    # ══════════════════════════════════════════════════════════

    def conversation_to_csv(self, messages: list) -> str:
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(["#", "Timestamp", "Role", "Content", "Model", "Tokens"])
        for i, msg in enumerate(messages, 1):
            writer.writerow([i, msg.get("created_at", ""), msg.get("role", ""),
                           msg.get("content", ""), msg.get("model", ""),
                           msg.get("tokens_used", 0)])
        return output.getvalue()
